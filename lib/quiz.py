import json
import random
import re
import secrets
import uuid
from html import unescape
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def default_answer(index: int) -> dict:
	return {
		"id": f"ans_{uuid.uuid4().hex[:8]}",
		"text": f"Option {index + 1}",
	}


def default_question(index: int) -> dict:
	return {
		"id": f"q_{uuid.uuid4().hex[:8]}",
		"title": f"New Question {index + 1}",
		"question_text": "<div><p>Your question here</p></div>",
		"question_type": "multiple_choice",
		"points": 1,
		"answers": [default_answer(i) for i in range(4)],
		"correct_answer_ids": [],
		"feedback": "<div><p>Explain the answer here.</p></div>",
	}


def default_quiz() -> dict:
	return {
		"assessment_id": f"quiz_{uuid.uuid4().hex[:8]}",
		"quiz_title": "New Quiz",
		"question_groups": [
			{
				"title": "Question Group 1",
				"questions_to_select": 1,
				"questions": [default_question(0)],
			}
		],
	}


def normalize_quiz(payload: dict) -> dict:
	normalized = {
		"assessment_id": str(payload.get("assessment_id") or f"quiz_{uuid.uuid4().hex[:8]}"),
		"quiz_title": str(payload.get("quiz_title") or "Untitled Quiz"),
		"question_groups": [],
	}

	groups = payload.get("question_groups")
	if not isinstance(groups, list):
		raise ValueError("'question_groups' must be a list.")

	for group_index, group in enumerate(groups):
		if not isinstance(group, dict):
			raise ValueError(f"Group at index {group_index} must be an object.")

		group_title = str(group.get("title") or f"Question Group {group_index + 1}")
		questions = group.get("questions")
		if not isinstance(questions, list):
			raise ValueError(f"'questions' in group {group_index + 1} must be a list.")

		normalized_group = {
			"title": group_title,
			"questions_to_select": None,
			"questions": [],
		}

		for question_index, question in enumerate(questions):
			if not isinstance(question, dict):
				raise ValueError(
					f"Question at index {question_index} in group {group_index + 1} must be an object."
				)

			answers = question.get("answers")
			if not isinstance(answers, list):
				raise ValueError(
					f"'answers' for question index {question_index} in group {group_index + 1} must be a list."
				)

			normalized_answers = []
			for answer_index, answer in enumerate(answers):
				if not isinstance(answer, dict):
					raise ValueError(
						f"Answer at index {answer_index} in question {question_index + 1} must be an object."
					)
				normalized_answers.append(
					{
						"id": str(answer.get("id") or f"ans_{uuid.uuid4().hex[:8]}"),
						"text": str(answer.get("text") or ""),
					}
				)

			raw_correct_ids = question.get("correct_answer_ids")
			if raw_correct_ids is None:
				single = question.get("correct_answer_id")
				if single:
					raw_correct_ids = [single]
				else:
					raw_correct_ids = []

			if not isinstance(raw_correct_ids, list):
				raise ValueError(
					f"'correct_answer_ids' for question index {question_index} in group {group_index + 1} must be a list."
				)

			answer_id_set = {a["id"] for a in normalized_answers}
			filtered_correct_ids = [
				str(correct_id)
				for correct_id in raw_correct_ids
				if str(correct_id) in answer_id_set
			]

			normalized_question = {
				"id": str(question.get("id") or f"q_{uuid.uuid4().hex[:8]}"),
				"title": str(question.get("title") or f"Question {question_index + 1}"),
				"question_text": str(question.get("question_text") or ""),
				"question_type": str(question.get("question_type") or "multiple_choice"),
				"points": int(question.get("points") or 0),
				"answers": normalized_answers,
				"correct_answer_ids": filtered_correct_ids,
				"feedback": str(question.get("feedback") or ""),
			}
			normalized_group["questions"].append(normalized_question)

		raw_questions_to_select = group.get("questions_to_select")
		if raw_questions_to_select is None:
			normalized_group["questions_to_select"] = len(normalized_group["questions"])
		else:
			try:
				normalized_group["questions_to_select"] = int(raw_questions_to_select)
			except (TypeError, ValueError) as exc:
				raise ValueError(
					f"'questions_to_select' in group {group_index + 1} must be a whole number."
				) from exc

		normalized["question_groups"].append(normalized_group)

	return normalized


def export_quiz_payload(quiz: dict) -> dict:
	exported = {
		"assessment_id": quiz.get("assessment_id", ""),
		"quiz_title": quiz.get("quiz_title", ""),
		"question_groups": [],
	}

	for group in quiz.get("question_groups", []):
		group_questions = group.get("questions", [])
		default_select_count = len(group_questions)
		raw_select_count = group.get("questions_to_select", default_select_count)
		try:
			select_count = int(raw_select_count)
		except (TypeError, ValueError):
			select_count = default_select_count

		out_group = {
			"title": group.get("title", ""),
			"questions_to_select": select_count,
			"questions": [],
		}

		for question in group_questions:
			out_question = {
				"id": question.get("id", ""),
				"title": question.get("title", ""),
				"question_text": question.get("question_text", ""),
				"question_type": question.get("question_type", "multiple_choice"),
				"points": int(question.get("points", 0)),
				"answers": question.get("answers", []),
				"correct_answer_ids": list(question.get("correct_answer_ids", [])),
				"feedback": question.get("feedback", ""),
			}

			if len(out_question["correct_answer_ids"]) == 1:
				out_question["correct_answer_id"] = out_question["correct_answer_ids"][0]

			out_group["questions"].append(out_question)

		exported["question_groups"].append(out_group)

	return exported


def load_json_text(json_text: str) -> dict:
	loaded = json.loads(json_text)
	return normalize_quiz(loaded)


def refresh_preview(quiz: dict) -> str:
	payload = export_quiz_payload(quiz)
	return json.dumps(payload, indent=4, ensure_ascii=False)


def save_quiz_to_path(preview_text: str, save_path: str) -> Path:
	path = Path(save_path).expanduser()
	path.parent.mkdir(parents=True, exist_ok=True)
	path.write_text(preview_text, encoding="utf-8")
	return path


def _html_to_text(raw_html: str) -> str:
	text = str(raw_html or "")
	text = re.sub(r"(?i)<br\\s*/?>", "\n", text)
	text = re.sub(r"(?i)</(p|div|h[1-6]|li|tr)>", "\n", text)
	text = re.sub(r"(?i)<li[^>]*>", "- ", text)
	text = re.sub(r"<[^>]+>", "", text)
	text = unescape(text)
	text = text.replace("\r\n", "\n").replace("\r", "\n")
	text = re.sub(r"\n{3,}", "\n\n", text)
	return text.strip()


def _flatten_questions(quiz_payload: dict) -> list[tuple[int, str, dict]]:
	flattened = []
	number = 1

	for group_index, group in enumerate(quiz_payload.get("question_groups", [])):
		group_title = str(group.get("title") or f"Question Group {group_index + 1}")
		for question in group.get("questions", []):
			flattened.append((number, group_title, question))
			number += 1

	return flattened


def _group_selection_count(group: dict, group_index: int) -> int:
	group_questions = group.get("questions", [])
	available = len(group_questions)
	raw_count = group.get("questions_to_select", available)

	try:
		count = int(raw_count)
	except (TypeError, ValueError) as exc:
		raise ValueError(
			f"Group {group_index + 1} selection count must be a whole number."
		) from exc

	if count < 1:
		raise ValueError(f"Group {group_index + 1} selection count must be at least 1.")
	if count > available:
		raise ValueError(
			f"Group {group_index + 1} requests {count} question(s) but only {available} are available."
		)

	return count


def _flatten_questions_for_permutation(quiz_payload: dict) -> list[tuple[int, str, dict]]:
	flattened: list[tuple[int, str, dict]] = []
	number = 1

	for group_index, group in enumerate(quiz_payload.get("question_groups", [])):
		group_title = str(group.get("title") or f"Question Group {group_index + 1}")
		group_questions = list(group.get("questions", []))
		pick_count = _group_selection_count(group, group_index)

		if pick_count == len(group_questions):
			selected_questions = group_questions
		else:
			picked_indexes = sorted(random.sample(range(len(group_questions)), pick_count))
			selected_questions = [group_questions[index] for index in picked_indexes]

		for question in selected_questions:
			flattened.append((number, group_title, question))
			number += 1

	return flattened


def generate_permutation_seed() -> int:
	return secrets.randbits(64)


def format_permutation_id(seed: int, permutation_number: int) -> str:
	return f"{int(seed):x}&{int(permutation_number):x}"


def parse_permutation_id(permutation_id: str) -> tuple[int, int]:
	parts = str(permutation_id).strip().split("&")
	if len(parts) != 2:
		raise ValueError("Invalid permutation ID format. Expected '<seed_hex>&<permutation_hex>'.")
	try:
		seed = int(parts[0], 16)
		permutation_number = int(parts[1], 16)
	except ValueError as exc:
		raise ValueError("Permutation ID contains non-hex values.") from exc
	if permutation_number < 1:
		raise ValueError("Permutation number in permutation ID must be >= 1.")
	return seed, permutation_number


def _flatten_questions_for_permutation_id(
	quiz_payload: dict,
	seed: int,
	permutation_number: int,
) -> list[tuple[int, str, dict]]:
	flattened: list[tuple[int, str, dict]] = []
	number = 1
	rng = random.Random(f"{int(seed):x}:{int(permutation_number):x}")

	for group_index, group in enumerate(quiz_payload.get("question_groups", [])):
		group_title = str(group.get("title") or f"Question Group {group_index + 1}")
		group_questions = list(group.get("questions", []))
		pick_count = _group_selection_count(group, group_index)

		if pick_count == len(group_questions):
			selected_questions = group_questions
		else:
			picked_indexes = sorted(rng.sample(range(len(group_questions)), pick_count))
			selected_questions = [group_questions[index] for index in picked_indexes]

		for question in selected_questions:
			flattened.append((number, group_title, question))
			number += 1

	return flattened


def _build_permutation_plan(quiz_payload: dict, permutations: int, seed: int | None = None) -> tuple[int, list[dict]]:
	permutation_count = max(1, int(permutations or 1))
	resolved_seed = int(seed if seed is not None else generate_permutation_seed())
	plans: list[dict] = []

	for permutation_number in range(1, permutation_count + 1):
		permutation_id = format_permutation_id(resolved_seed, permutation_number)
		flattened_questions = _flatten_questions_for_permutation_id(
			quiz_payload,
			resolved_seed,
			permutation_number,
		)
		plans.append(
			{
				"permutation_number": permutation_number,
				"permutation_id": permutation_id,
				"flattened_questions": flattened_questions,
			}
		)

	return resolved_seed, plans


def _greedy_paginate_questions(flattened_questions: list[tuple[int, str, dict]], questions_per_page: int) -> list[list[tuple[int, str, dict]]]:
	per_page = max(1, int(questions_per_page or 1))
	pages: list[list[tuple[int, str, dict]]] = []
	current_page: list[tuple[int, str, dict]] = []

	for question_entry in flattened_questions:
		if len(current_page) >= per_page:
			pages.append(current_page)
			current_page = []
		current_page.append(question_entry)

	if current_page:
		pages.append(current_page)

	return pages


def estimate_docx_sheet_count(quiz_payload: dict, questions_per_page: int, permutations: int = 1) -> int:
	permutation_count = max(1, int(permutations or 1))
	flattened_questions = _flatten_questions_for_permutation(quiz_payload)
	if not flattened_questions:
		return permutation_count
	pages_per_permutation = len(_greedy_paginate_questions(flattened_questions, questions_per_page))
	return pages_per_permutation * permutation_count


def _apply_orientation(document: Document, orientation: str) -> None:
	section = document.sections[0]
	normalized = str(orientation or "portrait").strip().lower()

	# Slightly tighter margins allow more content while remaining printable.
	section.top_margin = Inches(0.5)
	section.bottom_margin = Inches(0.5)
	section.left_margin = Inches(0.45)
	section.right_margin = Inches(0.45)

	cols_nodes = section._sectPr.xpath("./w:cols")
	if cols_nodes:
		cols = cols_nodes[0]
	else:
		cols = OxmlElement("w:cols")
		section._sectPr.append(cols)

	if normalized.startswith("land"):
		section.orientation = WD_ORIENT.LANDSCAPE
		if section.page_width < section.page_height:
			section.page_width, section.page_height = section.page_height, section.page_width
		cols.set(qn("w:num"), "3")
		cols.set(qn("w:space"), "240")
	else:
		section.orientation = WD_ORIENT.PORTRAIT
		if section.page_width > section.page_height:
			section.page_width, section.page_height = section.page_height, section.page_width
		cols.set(qn("w:num"), "1")


def build_docx_export(
	quiz_payload: dict,
	orientation: str,
	questions_per_page: int,
	permutations: int = 1,
	permutation_seed: int | None = None,
) -> bytes:
	document = Document()
	_apply_orientation(document, orientation)

	quiz_title = str(quiz_payload.get("quiz_title") or "Untitled Quiz")
	_, permutation_plans = _build_permutation_plan(
		quiz_payload,
		permutations=permutations,
		seed=permutation_seed,
	)

	for permutation_index, plan in enumerate(permutation_plans):
		if permutation_index > 0:
			document.add_page_break()

		document.add_heading(quiz_title, level=1)
		document.add_paragraph(f"Permutation ID: {plan['permutation_id']}")
		document.add_paragraph("Student ID: ____________________    Name: ____________________    Date: ____________________")
		document.add_paragraph("")

		flattened_questions = plan["flattened_questions"]
		pages = _greedy_paginate_questions(flattened_questions, questions_per_page)

		if not flattened_questions:
			document.add_paragraph("No questions available.")
			continue

		for page_index, page_questions in enumerate(pages):
			for question_number, _group_title, question in page_questions:
				points = int(question.get("points", 0) or 0)
				question_text = _html_to_text(question.get("question_text", ""))

				question_paragraph = document.add_paragraph()
				question_paragraph.add_run(f"{question_number}. ").bold = True
				if question_text:
					question_paragraph.add_run(question_text)
				if points > 0:
					question_paragraph.add_run(f" ({points} pts)")
				question_paragraph.paragraph_format.keep_with_next = True
				question_paragraph.paragraph_format.keep_together = True

				answer_paragraphs = []
				for answer_index, answer in enumerate(question.get("answers", [])):
					marker = chr(ord("A") + answer_index) if answer_index < 26 else str(answer_index + 1)
					answer_text = _html_to_text(answer.get("text", ""))
					answer_paragraph = document.add_paragraph(f"{marker}. {answer_text}")
					answer_paragraph.paragraph_format.space_before = Pt(0)
					answer_paragraph.paragraph_format.space_after = Pt(1)
					answer_paragraph.paragraph_format.keep_together = True
					answer_paragraphs.append(answer_paragraph)

				for answer_paragraph in answer_paragraphs[:-1]:
					answer_paragraph.paragraph_format.keep_with_next = True

				spacer = document.add_paragraph("")
				spacer.paragraph_format.space_before = Pt(0)
				spacer.paragraph_format.space_after = Pt(2)

			if page_index < len(pages) - 1:
				document.add_page_break()

	buffer = BytesIO()
	document.save(buffer)
	return buffer.getvalue()


def build_docx_answer_key_export(
	quiz_payload: dict,
	orientation: str,
	permutations: int = 1,
	permutation_seed: int | None = None,
) -> bytes:
	document = Document()
	_apply_orientation(document, orientation)

	quiz_title = str(quiz_payload.get("quiz_title") or "Untitled Quiz")
	_, permutation_plans = _build_permutation_plan(
		quiz_payload,
		permutations=permutations,
		seed=permutation_seed,
	)

	for permutation_index, plan in enumerate(permutation_plans):
		if permutation_index > 0:
			document.add_page_break()

		document.add_heading(f"{quiz_title} - Answer Key", level=1)
		document.add_paragraph(f"Permutation ID: {plan['permutation_id']}")
		document.add_paragraph("")

		flattened_questions = plan["flattened_questions"]
		if not flattened_questions:
			document.add_paragraph("No questions available.")
			continue

		for question_number, _group_title, question in flattened_questions:
			question_text = _html_to_text(question.get("question_text", "")) or f"Question {question_number}"
			document.add_paragraph(f"{question_number}. {question_text}")

			correct_ids = set(question.get("correct_answer_ids", []))
			if not correct_ids:
				document.add_paragraph("   Correct: (none)")
				continue

			correct_lines: list[str] = []
			for answer_index, answer in enumerate(question.get("answers", [])):
				answer_id = str(answer.get("id", ""))
				if answer_id not in correct_ids:
					continue
				marker = chr(ord("A") + answer_index) if answer_index < 26 else str(answer_index + 1)
				answer_text = _html_to_text(answer.get("text", ""))
				correct_lines.append(f"{marker}. {answer_text}")

			if correct_lines:
				document.add_paragraph("   Correct: " + " | ".join(correct_lines))
			else:
				document.add_paragraph("   Correct: (none)")

			document.add_paragraph("")

	buffer = BytesIO()
	document.save(buffer)
	return buffer.getvalue()
